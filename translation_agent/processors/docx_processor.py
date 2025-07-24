"""
DOCX document processor
"""
import copy
from typing import Dict, Any, List

import docx
from docx import Document

from .base import BaseDocumentProcessor

class DOCXProcessor(BaseDocumentProcessor):
    """Processor for DOCX documents"""
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text content from DOCX while preserving structure"""
        doc = Document(file_path)
        content = {
            'paragraphs': [],
            'tables': [],
            'headers': [],
            'footers': []
        }
        
        # Extract paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                content['paragraphs'].append({
                    'paragraph_index': para_idx,
                    'text': paragraph.text.strip(),
                    'original_text': paragraph.text.strip(),
                    'style': paragraph.style.name if paragraph.style else None
                })
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                'table_index': table_idx,
                'rows': []
            }
            
            for row_idx, row in enumerate(table.rows):
                row_data = {
                    'row_index': row_idx,
                    'cells': []
                }
                
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    row_data['cells'].append({
                        'cell_index': cell_idx,
                        'text': cell_text,
                        'original_text': cell_text
                    })
                
                table_data['rows'].append(row_data)
            
            content['tables'].append(table_data)
        
        # Extract headers and footers
        for section in doc.sections:
            if section.header:
                for para_idx, paragraph in enumerate(section.header.paragraphs):
                    if paragraph.text.strip():
                        content['headers'].append({
                            'paragraph_index': para_idx,
                            'text': paragraph.text.strip(),
                            'original_text': paragraph.text.strip()
                        })
            
            if section.footer:
                for para_idx, paragraph in enumerate(section.footer.paragraphs):
                    if paragraph.text.strip():
                        content['footers'].append({
                            'paragraph_index': para_idx,
                            'text': paragraph.text.strip(),
                            'original_text': paragraph.text.strip()
                        })
        
        return content
    
    def reconstruct_document(self, original_path: str, translated_content: Dict[str, Any], output_path: str) -> str:
        """Reconstruct DOCX by modifying original document"""
        doc = Document(original_path)
        
        # Update paragraphs
        for para_data in translated_content['paragraphs']:
            para_idx = para_data['paragraph_index']
            if para_idx < len(doc.paragraphs):
                doc.paragraphs[para_idx].text = para_data['text']
        
        # Update tables
        for table_data in translated_content['tables']:
            table_idx = table_data['table_index']
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                for row_data in table_data['rows']:
                    row_idx = row_data['row_index']
                    if row_idx < len(table.rows):
                        row = table.rows[row_idx]
                        for cell_data in row_data['cells']:
                            cell_idx = cell_data['cell_index']
                            if cell_idx < len(row.cells):
                                row.cells[cell_idx].text = cell_data['text']
        
        # Update headers and footers
        for section in doc.sections:
            if section.header:
                for para_data in translated_content['headers']:
                    para_idx = para_data['paragraph_index']
                    if para_idx < len(section.header.paragraphs):
                        section.header.paragraphs[para_idx].text = para_data['text']
            
            if section.footer:
                for para_data in translated_content['footers']:
                    para_idx = para_data['paragraph_index']
                    if para_idx < len(section.footer.paragraphs):
                        section.footer.paragraphs[para_idx].text = para_data['text']
        
        doc.save(output_path)
        return output_path
    
    def get_translatable_texts(self, extracted_content: Dict[str, Any]) -> List[str]:
        """Extract all translatable texts from DOCX content"""
        texts = []
        
        # Paragraphs
        for para_data in extracted_content['paragraphs']:
            texts.append(para_data['text'])
        
        # Tables
        for table_data in extracted_content['tables']:
            for row_data in table_data['rows']:
                for cell_data in row_data['cells']:
                    if cell_data['text'].strip():
                        texts.append(cell_data['text'])
        
        # Headers and footers
        for header_data in extracted_content['headers']:
            texts.append(header_data['text'])
        
        for footer_data in extracted_content['footers']:
            texts.append(footer_data['text'])
        
        return texts
    
    def apply_translations(self, extracted_content: Dict[str, Any], translations: List[str]) -> Dict[str, Any]:
        """Apply translations to DOCX content structure"""
        translated_content = copy.deepcopy(extracted_content)
        translation_idx = 0
        
        # Apply to paragraphs
        for para_data in translated_content['paragraphs']:
            if translation_idx < len(translations):
                para_data['text'] = translations[translation_idx]
                translation_idx += 1
        
        # Apply to tables
        for table_data in translated_content['tables']:
            for row_data in table_data['rows']:
                for cell_data in row_data['cells']:
                    if cell_data['text'].strip() and translation_idx < len(translations):
                        cell_data['text'] = translations[translation_idx]
                        translation_idx += 1
        
        # Apply to headers
        for header_data in translated_content['headers']:
            if translation_idx < len(translations):
                header_data['text'] = translations[translation_idx]
                translation_idx += 1
        
        # Apply to footers
        for footer_data in translated_content['footers']:
            if translation_idx < len(translations):
                footer_data['text'] = translations[translation_idx]
                translation_idx += 1
        
        return translated_content
